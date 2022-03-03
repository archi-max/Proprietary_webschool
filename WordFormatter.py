from docx.enum.text import WD_COLOR_INDEX
from pygments.formatter import Formatter
from enum import Enum
from docx import Document
from docx.shared import RGBColor

class WordStyle(Enum):
    """
    Enum for the different styles of words
    """
    COLOR = "color"
    BOLD = "bold"
    ITALIC = "italic"
    UNDERLINE = "underline"
    BGCOLOR = "bgcolor"

class WordFormatter(Formatter):

    def __init__(self, **options):
        Formatter.__init__(self, **options)

        # create a dict of (start, end) tuples that wrap the
        # value of a token so that we can use it in the format
        # method later
        self.styles = {}

        # we iterate over the `_styles` attribute of a style item
        # that contains the parsed style values.
        style_type = None
        for token, style in self.style:
            start = end = ''
            # a style item is a tuple in the following form:
            # colors are readily specified in hex: 'RRGGBB'
            if style['color']:
                style_type = style['color'].replace('#', '')
            if style['bold']:
                style_type = WordStyle.BOLD
            if style['italic']:
                style_type = WordStyle.ITALIC
            if style['underline']:
                style_type = WordStyle.UNDERLINE
            if style['bgcolor']:
                style_type = style['bgcolor'].replace('#', '')

            self.styles[token] = style_type

    def format(self, tokensource, outfile):
        # lastval is a string we use for caching
        # because it's possible that an lexer yields a number
        # of consecutive tokens with the same token type.
        # to minimize the size of the generated html markup we
        # try to join the values of same-type tokens here
        lastval = ''
        lasttype = None

        # wrap the whole output with <pre>
        document = Document()
        p = document.add_paragraph()

        for ttype, value in tokensource:
            # if the token type doesn't exist in the stylemap
            # we try it with the parent of the token type
            # eg: parent of Token.Literal.String.Double is
            # Token.Literal.String
            while ttype not in self.styles:
                if ttype is not None:
                    ttype = ttype.parent
                else:
                    print(value)
            if ttype == lasttype:
                # the current token type is the same of the last
                # iteration. cache it
                lastval += value
            else:
                # not the same token as last iteration, but we
                # have some data in the buffer. wrap it with the
                # defined style and write it to the output file
                if lastval:
                    if lasttype == 'color':
                        run = p.add_run(lastval)
                        run.font.color.rgb = RGBColor.from_string(self.styles[lasttype])
                    elif lasttype == WordStyle.BOLD:
                        p.add_run(lastval).bold = True
                    elif lasttype == WordStyle.ITALIC:
                        p.add_run(lastval).italic = True
                    elif lasttype == WordStyle.UNDERLINE:
                        p.add_run(lastval).underline = True
                    elif lasttype == WordStyle.BGCOLOR:
                        run = p.add_run(lastval)
                        run.font.highlight_color = WD_COLOR_INDEX.BLACK


                # set lastval/lasttype to current values
                lastval = value
                lasttype = ttype

        # if something is left in the buffer, write it to the
        # output file, then close the opened <pre> tag
        if lastval:
            if lasttype == 'color':
                run = p.add_run(lastval)
                run.font.color.rgb = RGBColor.from_string(self.styles[lasttype])
            elif lasttype == WordStyle.BOLD:
                p.add_run(lastval).bold = True
            elif lasttype == WordStyle.ITALIC:
                p.add_run(lastval).italic = True
            elif lasttype == WordStyle.UNDERLINE:
                p.add_run(lastval).underline = True
            elif lasttype == WordStyle.BGCOLOR:
                run = p.add_run(lastval)
                run.font.highlight_color = WD_COLOR_INDEX.BLACK
        document.save("testo.docx")